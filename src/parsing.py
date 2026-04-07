import os
import time
import logging
from pathlib import Path
from datetime import datetime, timezone

from schemas import (
    ParsedDocument, ExtractedTable, ExtractionError,
    generate_doc_id, generate_table_id, compute_table_confidence,
)

logger = logging.getLogger("extraction_pipeline")

PARSER_ROUTES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc_legacy",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xls",
    "text/csv": "csv",
    "text/plain": "text",
    "application/csv": "csv",
}

EXTENSION_OVERRIDES = {
    ".csv": "csv",
    ".tsv": "csv",
    ".xls": "xls",
    ".xlsx": "xlsx",
    ".doc": "doc_legacy",
    ".docx": "docx",
    ".pdf": "pdf",
    ".txt": "text",
    ".md": "text",
    ".log": "text",
    ".json": "text",
    ".xml": "text",
}


def detect_file_type(file_path):
    try:
        import magic
        mime_type = magic.from_file(file_path, mime=True)
    except ImportError:
        mime_type = _guess_mime_from_extension(file_path)
    except Exception:
        mime_type = _guess_mime_from_extension(file_path)

    extension = Path(file_path).suffix.lower()

    if mime_type == "text/plain" and extension in (".csv", ".tsv"):
        return mime_type, "csv"

    parser = PARSER_ROUTES.get(mime_type)
    if parser:
        return mime_type, parser

    parser = EXTENSION_OVERRIDES.get(extension)
    if parser:
        return mime_type, parser

    return mime_type, "unsupported"


def _guess_mime_from_extension(file_path):
    ext = Path(file_path).suffix.lower()
    mime_map = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".csv": "text/csv",
        ".tsv": "text/csv",
        ".txt": "text/plain",
        ".md": "text/plain",
        ".log": "text/plain",
        ".json": "text/plain",
        ".xml": "text/plain",
    }
    return mime_map.get(ext, "application/octet-stream")


def _now_iso():
    return datetime.now(timezone.utc).isoformat()


def _failed_document(file_path, doc_id, file_type, extension, exception, duration_ms):
    file_size = 0
    if os.path.exists(file_path):
        file_size = os.path.getsize(file_path)

    return ParsedDocument(
        doc_id=doc_id,
        source_path=file_path,
        file_name=os.path.basename(file_path),
        file_type=file_type,
        file_extension=extension,
        file_size_bytes=file_size,
        extracted_text="",
        pages=[],
        page_count=0,
        tables=[],
        table_count=0,
        metadata={},
        extraction_timestamp=_now_iso(),
        extraction_duration_ms=duration_ms,
        extraction_method="failed",
        extraction_quality_score=0.0,
        ocr_applied=False,
        encoding_detected=None,
        errors=[ExtractionError(
            stage="parsing",
            severity="error",
            message=f"Complete extraction failure: {str(exception)}",
            exception_type=type(exception).__name__,
        )],
        is_successful=False,
    )


def parse_pdf(file_path, doc_id=None):
    doc_id = doc_id or generate_doc_id()
    start_time = time.monotonic()
    errors = []
    pages = []
    tables = []
    metadata = {}
    ocr_applied = False

    try:
        import fitz

        mu_doc = fitz.open(file_path)
        page_count = len(mu_doc)

        pdf_meta = mu_doc.metadata or {}
        metadata = {
            "author": pdf_meta.get("author", ""),
            "title": pdf_meta.get("title", ""),
            "subject": pdf_meta.get("subject", ""),
            "creator": pdf_meta.get("creator", ""),
            "producer": pdf_meta.get("producer", ""),
            "creation_date": pdf_meta.get("creationDate", ""),
            "modification_date": pdf_meta.get("modDate", ""),
        }

        for page_num in range(page_count):
            page = mu_doc[page_num]
            try:
                text = page.get_text("text")
                pages.append({"page_number": page_num + 1, "text": text})
            except Exception as exc:
                errors.append(ExtractionError(
                    stage="parsing",
                    severity="warning",
                    message=f"Failed to extract text from page {page_num + 1}: {str(exc)}",
                    exception_type=type(exc).__name__,
                    page_number=page_num + 1,
                ))
                pages.append({"page_number": page_num + 1, "text": ""})

        mu_doc.close()

        total_text = "".join(p["text"] for p in pages)
        chars_per_page = len(total_text.strip()) / max(page_count, 1)

        if chars_per_page < 50:
            ocr_applied = _attempt_ocr(file_path, pages, errors)
            if ocr_applied:
                total_text = "".join(p["text"] for p in pages)

        tables = _extract_pdf_tables(file_path, errors)

        duration_ms = int((time.monotonic() - start_time) * 1000)

        return ParsedDocument(
            doc_id=doc_id,
            source_path=file_path,
            file_name=os.path.basename(file_path),
            file_type="application/pdf",
            file_extension=".pdf",
            file_size_bytes=os.path.getsize(file_path),
            extracted_text=total_text,
            pages=pages,
            page_count=page_count,
            tables=tables,
            table_count=len(tables),
            metadata=metadata,
            extraction_timestamp=_now_iso(),
            extraction_duration_ms=duration_ms,
            extraction_method="pymupdf+pdfplumber" + ("+tesseract" if ocr_applied else ""),
            extraction_quality_score=0.0,
            ocr_applied=ocr_applied,
            encoding_detected=None,
            errors=errors,
            is_successful=len(total_text.strip()) > 0,
        )

    except Exception as exc:
        duration_ms = int((time.monotonic() - start_time) * 1000)
        return _failed_document(file_path, doc_id, "application/pdf", ".pdf", exc, duration_ms)


def _extract_pdf_tables(file_path, errors):
    tables = []
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as plumber_doc:
            for page_num, plumber_page in enumerate(plumber_doc.pages):
                try:
                    page_tables = plumber_page.extract_tables({
                        "vertical_strategy": "lines_strict",
                        "horizontal_strategy": "lines_strict",
                        "snap_tolerance": 4,
                        "join_tolerance": 4,
                        "min_words_vertical": 3,
                        "min_words_horizontal": 1,
                    })
                    for raw_table in (page_tables or []):
                        if not raw_table or len(raw_table) < 2:
                            continue
                        headers = [str(c or "").strip() for c in raw_table[0]]
                        rows = [
                            [str(c or "").strip() for c in row]
                            for row in raw_table[1:]
                        ]
                        tables.append(ExtractedTable(
                            table_id=generate_table_id(),
                            page_number=page_num + 1,
                            sheet_name=None,
                            headers=headers,
                            rows=rows,
                            row_count=len(rows),
                            col_count=len(headers),
                            extraction_method="pdfplumber",
                            confidence=compute_table_confidence(headers, rows),
                        ))
                except Exception as exc:
                    errors.append(ExtractionError(
                        stage="table_extraction",
                        severity="warning",
                        message=f"Table extraction failed on page {page_num + 1}: {str(exc)}",
                        exception_type=type(exc).__name__,
                        page_number=page_num + 1,
                    ))
    except ImportError:
        errors.append(ExtractionError(
            stage="table_extraction",
            severity="warning",
            message="pdfplumber not installed, table extraction skipped",
        ))
    except Exception as exc:
        errors.append(ExtractionError(
            stage="table_extraction",
            severity="warning",
            message=f"pdfplumber failed to open file: {str(exc)}",
            exception_type=type(exc).__name__,
        ))

    return tables


def _attempt_ocr(file_path, pages, errors):
    try:
        import pytesseract
        from PIL import Image
        import io
        import fitz

        mu_doc = fitz.open(file_path)
        for page_num in range(len(mu_doc)):
            page = mu_doc[page_num]
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            ocr_text = pytesseract.image_to_string(img, lang="eng")
            if ocr_text.strip():
                pages[page_num]["text"] = ocr_text
        mu_doc.close()
        return True
    except ImportError:
        errors.append(ExtractionError(
            stage="ocr",
            severity="warning",
            message="pytesseract or tesseract binary not available, OCR skipped",
        ))
        return False
    except Exception as exc:
        errors.append(ExtractionError(
            stage="ocr",
            severity="error",
            message=f"OCR failed: {str(exc)}",
            exception_type=type(exc).__name__,
        ))
        return False


def parse_docx(file_path, doc_id=None):
    doc_id = doc_id or generate_doc_id()
    start_time = time.monotonic()
    errors = []
    tables = []
    metadata = {}

    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)

        core = doc.core_properties
        metadata = {
            "author": core.author or "",
            "title": core.title or "",
            "subject": core.subject or "",
            "created": str(core.created) if core.created else "",
            "modified": str(core.modified) if core.modified else "",
            "last_modified_by": core.last_modified_by or "",
        }

        paragraphs_text = [para.text for para in doc.paragraphs]
        full_text = "\n".join(paragraphs_text)
        pages = [{"page_number": 1, "text": full_text}]

        for tbl_idx, table in enumerate(doc.tables):
            try:
                rows_data = []
                for row in table.rows:
                    rows_data.append([cell.text.strip() for cell in row.cells])
                if len(rows_data) >= 2:
                    headers = rows_data[0]
                    data_rows = rows_data[1:]
                    tables.append(ExtractedTable(
                        table_id=generate_table_id(),
                        page_number=None,
                        sheet_name=None,
                        headers=headers,
                        rows=data_rows,
                        row_count=len(data_rows),
                        col_count=len(headers),
                        extraction_method="python-docx",
                        confidence=compute_table_confidence(headers, data_rows),
                    ))
            except Exception as exc:
                errors.append(ExtractionError(
                    stage="table_extraction",
                    severity="warning",
                    message=f"Table {tbl_idx} extraction failed: {str(exc)}",
                    exception_type=type(exc).__name__,
                ))

        header_footer_text = []
        for section in doc.sections:
            for hdr in [section.header, section.first_page_header]:
                if hdr and hdr.paragraphs:
                    for p in hdr.paragraphs:
                        if p.text.strip():
                            header_footer_text.append(p.text)
            for ftr in [section.footer, section.first_page_footer]:
                if ftr and ftr.paragraphs:
                    for p in ftr.paragraphs:
                        if p.text.strip():
                            header_footer_text.append(p.text)

        if header_footer_text:
            metadata["headers_footers"] = " | ".join(header_footer_text)

        duration_ms = int((time.monotonic() - start_time) * 1000)

        return ParsedDocument(
            doc_id=doc_id,
            source_path=file_path,
            file_name=os.path.basename(file_path),
            file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_extension=".docx",
            file_size_bytes=os.path.getsize(file_path),
            extracted_text=full_text,
            pages=pages,
            page_count=1,
            tables=tables,
            table_count=len(tables),
            metadata=metadata,
            extraction_timestamp=_now_iso(),
            extraction_duration_ms=duration_ms,
            extraction_method="python-docx",
            extraction_quality_score=0.0,
            ocr_applied=False,
            encoding_detected=None,
            errors=errors,
            is_successful=len(full_text.strip()) > 0,
        )

    except Exception as exc:
        duration_ms = int((time.monotonic() - start_time) * 1000)
        return _failed_document(
            file_path, doc_id,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".docx", exc, duration_ms,
        )


def parse_doc_legacy(file_path, doc_id=None):
    doc_id = doc_id or generate_doc_id()
    start_time = time.monotonic()
    errors = []

    import subprocess
    import tempfile

    try:
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0 and result.stdout.strip():
            text = result.stdout
            duration_ms = int((time.monotonic() - start_time) * 1000)
            return ParsedDocument(
                doc_id=doc_id,
                source_path=file_path,
                file_name=os.path.basename(file_path),
                file_type="application/msword",
                file_extension=".doc",
                file_size_bytes=os.path.getsize(file_path),
                extracted_text=text,
                pages=[{"page_number": 1, "text": text}],
                page_count=1,
                tables=[],
                table_count=0,
                metadata={"conversion_method": "antiword"},
                extraction_timestamp=_now_iso(),
                extraction_duration_ms=duration_ms,
                extraction_method="antiword",
                extraction_quality_score=0.0,
                ocr_applied=False,
                encoding_detected=None,
                errors=errors,
                is_successful=True,
            )
    except FileNotFoundError:
        errors.append(ExtractionError(
            stage="parsing", severity="warning",
            message="antiword not installed, falling back to LibreOffice conversion",
        ))
    except Exception:
        pass

    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx",
                 "--outdir", tmp_dir, file_path],
                capture_output=True, timeout=120, check=True,
            )
            docx_name = Path(file_path).stem + ".docx"
            converted_path = os.path.join(tmp_dir, docx_name)
            result = parse_docx(converted_path, doc_id)
            result.source_path = file_path
            result.file_name = os.path.basename(file_path)
            result.file_type = "application/msword"
            result.file_extension = ".doc"
            result.extraction_method = "libreoffice+python-docx"
            result.metadata["conversion_method"] = "libreoffice"
            result.errors.extend(errors)
            return result
    except Exception as exc:
        duration_ms = int((time.monotonic() - start_time) * 1000)
        return _failed_document(file_path, doc_id, "application/msword", ".doc", exc, duration_ms)


def parse_spreadsheet(file_path, doc_id=None, parser_hint="xlsx"):
    doc_id = doc_id or generate_doc_id()
    start_time = time.monotonic()
    errors = []
    tables = []
    metadata = {}
    extension = Path(file_path).suffix.lower()
    all_text_parts = []

    try:
        import pandas as pd

        if parser_hint == "csv":
            encoding, encoding_confidence = _detect_encoding(file_path)
            metadata["detected_encoding"] = encoding
            metadata["encoding_confidence"] = str(encoding_confidence)

            delimiter = _detect_csv_delimiter(file_path, encoding)
            metadata["delimiter"] = delimiter

            df = pd.read_csv(
                file_path, encoding=encoding, delimiter=delimiter,
                on_bad_lines="warn", engine="python", dtype=str,
            )
            sheet_map = {"Sheet1": df}

        elif parser_hint == "xlsx":
            import openpyxl
            sheet_map = pd.read_excel(file_path, sheet_name=None, engine="openpyxl", dtype=str)
            try:
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                props = wb.properties
                metadata["author"] = props.creator or ""
                metadata["title"] = props.title or ""
                metadata["created"] = str(props.created) if props.created else ""
                metadata["modified"] = str(props.modified) if props.modified else ""
                metadata["sheet_names"] = ",".join(wb.sheetnames)
                metadata["sheet_count"] = str(len(wb.sheetnames))
                wb.close()
            except Exception as exc:
                errors.append(ExtractionError(
                    stage="parsing", severity="warning",
                    message=f"Failed to extract XLSX metadata: {str(exc)}",
                    exception_type=type(exc).__name__,
                ))

        elif parser_hint == "xls":
            sheet_map = pd.read_excel(file_path, sheet_name=None, engine="xlrd", dtype=str)
            try:
                import xlrd
                wb = xlrd.open_workbook(file_path)
                metadata["sheet_names"] = ",".join(wb.sheet_names())
                metadata["sheet_count"] = str(wb.nsheets)
            except Exception:
                pass

        else:
            raise ValueError(f"Unknown parser_hint: {parser_hint}")

        for sheet_name, df in sheet_map.items():
            df = df.fillna("")
            headers = [str(c) for c in df.columns.tolist()]
            rows = [[str(v) for v in row] for row in df.values.tolist()]

            tables.append(ExtractedTable(
                table_id=generate_table_id(),
                page_number=None,
                sheet_name=str(sheet_name),
                headers=headers,
                rows=rows,
                row_count=len(rows),
                col_count=len(headers),
                extraction_method=f"pandas+{parser_hint}",
                confidence=1.0,
            ))

            text_repr = f"[Sheet: {sheet_name}]\n"
            text_repr += "\t".join(headers) + "\n"
            for row in rows[:500]:
                text_repr += "\t".join(row) + "\n"
            if len(rows) > 500:
                text_repr += f"... ({len(rows) - 500} additional rows truncated)\n"
            all_text_parts.append(text_repr)

        full_text = "\n\n".join(all_text_parts)

        mime_map = {
            "csv": "text/csv",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "xls": "application/vnd.ms-excel",
        }

        duration_ms = int((time.monotonic() - start_time) * 1000)

        return ParsedDocument(
            doc_id=doc_id,
            source_path=file_path,
            file_name=os.path.basename(file_path),
            file_type=mime_map.get(parser_hint, "application/octet-stream"),
            file_extension=extension,
            file_size_bytes=os.path.getsize(file_path),
            extracted_text=full_text,
            pages=[{"page_number": 1, "text": full_text}],
            page_count=1,
            tables=tables,
            table_count=len(tables),
            metadata=metadata,
            extraction_timestamp=_now_iso(),
            extraction_duration_ms=duration_ms,
            extraction_method=f"pandas+{parser_hint}",
            extraction_quality_score=0.0,
            ocr_applied=False,
            encoding_detected=metadata.get("detected_encoding"),
            errors=errors,
            is_successful=len(full_text.strip()) > 0,
        )

    except Exception as exc:
        duration_ms = int((time.monotonic() - start_time) * 1000)
        mime_map = {
            "csv": "text/csv",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "xls": "application/vnd.ms-excel",
        }
        return _failed_document(
            file_path, doc_id,
            mime_map.get(parser_hint, "application/octet-stream"),
            extension, exc, duration_ms,
        )


def parse_text(file_path, doc_id=None):
    doc_id = doc_id or generate_doc_id()
    start_time = time.monotonic()
    errors = []

    encoding, confidence = _detect_encoding(file_path)

    try:
        with open(file_path, "r", encoding=encoding, errors="replace") as fh:
            text = fh.read()
    except Exception as exc:
        errors.append(ExtractionError(
            stage="parsing", severity="warning",
            message=f"Encoding {encoding} failed, falling back to latin-1: {str(exc)}",
            exception_type=type(exc).__name__,
        ))
        with open(file_path, "r", encoding="latin-1") as fh:
            text = fh.read()
        encoding = "latin-1"

    duration_ms = int((time.monotonic() - start_time) * 1000)

    return ParsedDocument(
        doc_id=doc_id,
        source_path=file_path,
        file_name=os.path.basename(file_path),
        file_type="text/plain",
        file_extension=Path(file_path).suffix.lower(),
        file_size_bytes=os.path.getsize(file_path),
        extracted_text=text,
        pages=[{"page_number": 1, "text": text}],
        page_count=1,
        tables=[],
        table_count=0,
        metadata={
            "detected_encoding": encoding,
            "encoding_confidence": str(confidence),
            "line_count": str(text.count("\n") + 1),
            "char_count": str(len(text)),
        },
        extraction_timestamp=_now_iso(),
        extraction_duration_ms=duration_ms,
        extraction_method="direct_read+chardet",
        extraction_quality_score=0.0,
        ocr_applied=False,
        encoding_detected=encoding,
        errors=errors,
        is_successful=len(text.strip()) > 0,
    )


def _detect_encoding(file_path):
    try:
        import chardet
        with open(file_path, "rb") as fh:
            raw = fh.read(65536)
            detection = chardet.detect(raw)
            encoding = detection["encoding"] or "utf-8"
            confidence = detection.get("confidence", 0)
            return encoding, confidence
    except ImportError:
        return "utf-8", 1.0


def _detect_csv_delimiter(file_path, encoding):
    import csv
    try:
        with open(file_path, "r", encoding=encoding, errors="replace") as fh:
            sample = fh.read(8192)
            dialect = csv.Sniffer().sniff(sample)
            return dialect.delimiter
    except Exception:
        return ","


def compute_quality_score(doc):
    if not doc.is_successful:
        return 0.0

    chars_per_page = len(doc.extracted_text) / max(doc.page_count, 1)
    if chars_per_page < 10:
        text_density = 0.0
    elif chars_per_page < 200:
        text_density = chars_per_page / 200.0
    elif chars_per_page <= 5000:
        text_density = 1.0
    else:
        text_density = max(0.5, 1.0 - (chars_per_page - 5000) / 50000)

    error_count = len([e for e in doc.errors if e.severity == "error"])
    warning_count = len([e for e in doc.errors if e.severity == "warning"])
    error_penalty = max(0.0, 1.0 - (error_count * 0.3 + warning_count * 0.1))

    if doc.tables:
        table_integrity = sum(t.confidence for t in doc.tables) / len(doc.tables)
    else:
        table_integrity = 1.0

    key_fields = ["author", "title", "created", "creation_date"]
    populated = sum(1 for k in key_fields if doc.metadata.get(k, "").strip())
    metadata_completeness = populated / max(len(key_fields), 1)

    enc_conf_str = doc.metadata.get("encoding_confidence", "1.0")
    try:
        encoding_confidence = float(enc_conf_str)
    except ValueError:
        encoding_confidence = 1.0

    score = (
        text_density * 0.4
        + error_penalty * 0.2
        + table_integrity * 0.2
        + metadata_completeness * 0.1
        + encoding_confidence * 0.1
    )

    return round(min(max(score, 0.0), 1.0), 3)


def log_extraction_result(doc):
    if doc.is_successful:
        if doc.errors:
            logger.warning(
                f"PARTIAL_EXTRACTION doc_id={doc.doc_id} file={doc.file_name} "
                f"method={doc.extraction_method} quality={doc.extraction_quality_score:.3f} "
                f"warnings={len(doc.errors)} duration_ms={doc.extraction_duration_ms}"
            )
        else:
            logger.info(
                f"EXTRACTION_SUCCESS doc_id={doc.doc_id} file={doc.file_name} "
                f"method={doc.extraction_method} quality={doc.extraction_quality_score:.3f} "
                f"pages={doc.page_count} tables={doc.table_count} "
                f"duration_ms={doc.extraction_duration_ms}"
            )
    else:
        logger.error(
            f"EXTRACTION_FAILED doc_id={doc.doc_id} file={doc.file_name} "
            f"method={doc.extraction_method} errors={[e.message for e in doc.errors]} "
            f"duration_ms={doc.extraction_duration_ms}"
        )
